import os
import re
from PIL import Image
from google import genai

MODEL_ID = "gemini-3.5-flash"


def _load_env_file(env_path=".env"):
    if not os.path.exists(env_path):
        return

    with open(env_path, "r", encoding="utf-8") as env_file:
        for raw_line in env_file:
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue

            key, value = line.split("=", 1)
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if key and key not in os.environ:
                os.environ[key] = value


_load_env_file()
API_KEY = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or ""


def create_client():
    api_key = API_KEY.strip() if API_KEY else ""
    if api_key:
        return genai.Client(api_key=api_key)
    print("⚠️ Warning: API key is not set. Set GEMINI_API_KEY or GOOGLE_API_KEY in your environment.")
    return genai.Client()


client = create_client()

# Folder structure paths
def get_exam_folders(grade, semester):
    """Generate folder paths based on grade and semester."""
    folder_prefix = f"{grade}ap-{semester}"
    return {
        "exam1": f"EXAMS/{folder_prefix}-exam1",
        "exam2": f"EXAMS/{folder_prefix}-exam2",
    }

def get_images_from_folder(folder_path):
    """Load all .png/.jpg/.jpeg files from a folder (sorted)."""
    if not os.path.exists(folder_path):
        print(f"❌ Error: Folder {folder_path} does not exist.")
        return []
    return [
        os.path.join(folder_path, f)
        for f in sorted(os.listdir(folder_path))
        if f.lower().endswith(('.png', '.jpg', '.jpeg'))
    ]

def generate_initial_exam(user_theme, school_name, town_name, grade, semester, exam_folders):
    """Generate the initial exam based on user theme, school name, town name, grade, and semester."""
    # Load example images for few-shot
    exam1_images = get_images_from_folder(exam_folders["exam1"])

    # Verify examples exist
    if not exam1_images:
        print(f"❌ Error: Missing required examples for Grade {grade} Semester {semester} Exam 1!")
        return None


    grade_specific_requirements = ""

    if grade == "2":
        grade_specific_requirements = f"""
        - Do not include any placeholders or mentions of logos.
        - Do not include or refer to any dates.
        - Ensure that any words students must find (e.g., synonyms or antonyms) appear verbatim in the passage.
        - You must strictly follow the exact structure provided below.
        - Provide all headings in Arabic.
        - Group related grammar questions together in single items.
        - Never put similar consecutive questions, combine them.
        - This exam is for الصف {grade} الابتدائي، لذا لا تقم بطرح أي سؤال عن الإعراب.
        """
    elif grade == "3":
        grade_specific_requirements = f"""
        - Do not include any placeholders or mentions of logos (e.g., remove flag/logos).
        - Do not include or refer to any dates.
        - Ensure that any words students must find (e.g., synonyms or antonyms) appear verbatim in the passage.
        - You must strictly follow the exact structure provided below, without any deviations.
        - Provide all headings in Arabic.
        - This exam is for الصف الثالث الابتدائي، لذا لا تقم بطرح أي سؤال عن الإعراب.
        - For grammar (البناء اللغوي), group related tasks into as few questions as possible by including multiple items in one.
        - Do NOT repeat the same question type consecutively; if combining items, place them in the same question.
        - Model answers should be fully in Arabic.
        - Adapt the difficulty level appropriately for grade {grade} students.
        """
    elif grade == "4":
        grade_specific_requirements = f"""
        - Do not include any placeholders or mentions of logos (e.g., remove flag/logos).
        - Do not include or refer to any dates.
        - Ensure that any words students must find (e.g., synonyms or antonyms) appear verbatim in the passage.
        - You must strictly follow the exact structure provided below, without any deviations.
        - Provide all headings in Arabic.
        - For grammar (البناء اللغوي), group related tasks into as few questions as possible by including multiple items in one.
        - Do NOT repeat the same question type consecutively; if combining items, place them in the same question.
        - For any grammar-parsing question (إعراب), the word to parse must appear in the reading passage enclosed in parentheses “(…)”, and the question must be phrased exactly as: “أعرب ما بين قوسين في النص”.
        - Model answers should be fully in Arabic.
        - Adapt the difficulty level appropriately for grade {grade} students.
        """
    elif grade == "5":
        grade_specific_requirements = f"""
        - Do not include any placeholders or mentions of logos (e.g., remove flag/logos).
        - Do not include or refer to any dates.
        - Ensure that any words students must find (e.g., synonyms or antonyms) appear verbatim in the passage.
        - You must strictly follow the exact structure provided below, without any deviations.
        - Provide all headings in Arabic.
        - For grammar (البناء اللغوي), group related tasks into as few questions as possible by including multiple items in one.
        - Do NOT repeat the same question type consecutively; if combining items, place them in the same question.
        - For any grammar-parsing question (إعراب), the word to parse must appear in the reading passage enclosed in parentheses “(…)”، و السؤال يجب أن يكون بصيغة: “أعرب ما بين قوسين في النص”.
        - Model answers should be fully in Arabic.
        - Adapt the difficulty level appropriately for grade {grade} students.
        """

    # The prompt include grade and semester information
    prompt = [
        f"""
        You are an expert Arabic language teacher.
        Create an original Arabic exam and answer key that exactly mirror the provided examples.

        Grade Level: {grade} (الصف {grade} الابتدائي)
        Semester: {semester} (الفصل {semester})
        Theme: {user_theme}

        Requirements:
        {grade_specific_requirements}

        VERY IMPORTANT:
        - You MUST use exactly the school name '{school_name}' for مدرسة
        - You MUST use exactly the town name '{town_name}' for مديرية التربية لولاية
        - Do NOT replace these with placeholders or other values.

        VERY IMPORTANT CLARIFICATION:
        Questions about finding synonyms (مرادف) and antonyms (ضد) should be placed under البناء الفكري (conceptual construction) section, NOT under البناء اللغوي (linguistic construction).

        EXACT STRUCTURE FOR THE EXAM:

        الجمهورية الجزائرية الديمقراطية الشعبية
        وزارة التربية الوطنية
        مديرية التربية لولاية {town_name}
        مدرسة {school_name}
        اختبار الفصل {semester} في مادة اللغة العربية

        النص
        [Insert appropriate reading passage here related to the theme]

        1. البناء الفكري
        1.1 [General question about the text]
        1.2 [General question about the text]
        1.3 [Question about finding a synonym: استخرج من النص مرادف كلمة "..."]
        1.4 [Question about finding an antonym: استخرج من النص ضد كلمة "..."]
        [Additional questions as needed]

        2. البناء اللغوي
        2.1 [Grammar question]
        2.2 [Grammar question]
        2.3 [Grammar question]
        2.4 [Grammar question]
        [Additional grammar questions as needed]

        الوضعية الادماجية
        [Insert integration situation/writing prompt here]

        EXACT STRUCTURE FOR THE CORRECTION:

        1. البناء الفكري
        1.1 [Answer]
        1.2 [Answer]
        1.3 [Answer to synonym question]
        1.4 [Answer to antonym question]
        [Additional answers as needed]

        2. البناء اللغوي
        2.1 [Answer to grammar question]
        2.2 [Answer to grammar question]
        2.3 [Answer to grammar question]
        2.4 [Answer to grammar question]
        [Additional answers as needed]

        الوضعية الادماجية
        [Model answer for the integration situation]

               TASK:
        Generate:
        1. A brand-new Arabic exam for grade {grade}, semester {semester} following EXACTLY the structure specified above, centered on the specified theme.
        2. A full correction sheet with model answers following EXACTLY the structure specified above.

        CRUCIAL REMINDER:
        - Questions about finding synonyms (مرادف) and antonyms (ضد) must be placed under البناء الفكري section, NOT under البناء اللغوي.
        - You MUST use the school name '{school_name}' for مدرسة
        - You MUST use the town name '{town_name}' for مديرية التربية لولاية
        - DO NOT replace these with placeholders like [PLACEHOLDER].
        - Create content appropriate for grade {grade} students in semester {semester}.

        YOU MUST ADHERE TO THE EXACT STRUCTURE PROVIDED - DO NOT MODIFY OR ADD SECTIONS.
        DONT ADD ANY TEXT OUTSIDE OF THE EXAM GENERATED IN YOUR RESPONSE.

        Clearly label the exam part with "**EXAM:**" and the correction part with "**CORRECTION:**"
        """
    ]

    def append_images(label, images):
        prompt.append(f"### {label}")
        for img_path in images:
            prompt.append(Image.open(img_path))

    # Append few-shot examples
    append_images(f"Example 1 – Grade {grade} Semester {semester} Exam (نموذج الامتحان 1)", exam1_images)



    # Final TASK instruction with emphasis on strict structure and header requirements

    try:
        response = client.models.generate_content(
            model=MODEL_ID,
            contents=prompt,
        )
        response_text = response.text

        # Save the generated exam to a file
        output_file = f"generated_arabic_exam_grade{grade}_semester{semester}.txt"
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(response_text)
        print(f"✅ Exam saved to '{output_file}'")

        return response_text
    except Exception as e:
        print(f"❌ Generation failed: {str(e)}")
        return None

# The regenerate_specific_question function accept full exam and question id and grade and semester info
def regenerate_specific_question(full_exam_text, question_id, reading_passage, user_requirements="", grade=None, semester=None):
    """Regenerate a specific question in the exam based on user requirements."""
    # Split the exam and correction parts
    parts = full_exam_text.split("**CORRECTION:**")
    if len(parts) != 2:
        print("❌ Error: Could not identify exam and correction parts.")
        return full_exam_text

    exam_part = parts[0].replace("**EXAM:**", "").strip()
    correction_part = parts[1].strip()

    if question_id.lower() == "idmajiya":
        # Handle integration situation (الوضعية الادماجية)
        prompt = f"""
        You are an expert Arabic language teacher chatbot. I have an Arabic exam with a reading passage about a specific theme.
        I want you to regenerate only the الوضعية الادماجية (integration situation) section.

        This exam is for grade {grade}, semester {semester}. Adjust difficulty accordingly.

        Here is the reading passage from the exam:

        {reading_passage}

        CURRENT FULL EXAM:
        {exam_part}

        CURRENT CORRECTION:
        {correction_part}

        USER'S SPECIFIC REQUIREMENTS FOR THE NEW الوضعية الادماجية:
        {user_requirements if user_requirements else "Create a different integration situation that fits well with the theme of the reading passage."}

        TASK:
        1. Generate ONLY a new الوضعية الادماجية (integration situation) that meets the user's requirements.
        2. Generate ONLY the matching model answer for the correction sheet.
        3. Do not change any other parts of the exam or correction.
        4. Maintain the appropriate difficulty level and style consistent with the rest of the exam.
        5. Make sure the content is appropriate for students in grade {grade if grade else '3'}.

        FORMAT YOUR RESPONSE EXACTLY LIKE THIS:

        NEW_IDMAJIYA: [Your new integration situation here]

        NEW_IDMAJIYA_ANSWER: [Your new model answer here]
        """

        try:
            response = client.models.generate_content(
                model=MODEL_ID,
                contents=prompt,
            )
            response_text = response.text

            # Extract new content
            idmajiya_match = re.search(r"NEW_IDMAJIYA:(.*?)(?=NEW_IDMAJIYA_ANSWER:|$)", response_text, re.DOTALL)
            idmajiya_answer_match = re.search(r"NEW_IDMAJIYA_ANSWER:(.*?)(?=$)", response_text, re.DOTALL)

            if not idmajiya_match or not idmajiya_answer_match:
                print("❌ Error: Could not extract new integration situation or its answer.")
                return full_exam_text

            new_idmajiya = idmajiya_match.group(1).strip()
            new_idmajiya_answer = idmajiya_answer_match.group(1).strip()

            # Update exam content without printing changes
            exam_pattern = r"(الوضعية الادماجية\s*\n)(.*?)(?=\s*\n\s*\\*\\*CORRECTION:\\*\\*|\s*$)"
            corr_pattern = r"(الوضعية الادماجية\s*\n)(.*?)(?=\s*$)"

            updated_exam_part = re.sub(exam_pattern, f"\\1\n{new_idmajiya}\n", exam_part, flags=re.DOTALL)
            updated_correction_part = re.sub(corr_pattern, f"\\1\n{new_idmajiya_answer}\n", correction_part, flags=re.DOTALL)
            updated_full_exam = f"**EXAM:**\n{updated_exam_part}\n\n**CORRECTION:**\n{updated_correction_part}"

            # Save updated exam and show only success message
            with open("updated_arabic_exam.txt", "w", encoding="utf-8") as f:
                f.write(updated_full_exam)
            print("✅ Updated exam saved to 'updated_arabic_exam.txt'")

            return updated_full_exam

        except Exception as e:
            print(f"❌ Regeneration failed: {e}")
            return full_exam_text

    else:
        # Handle regular questions (البناء الفكري and البناء اللغوي)
        section_number = question_id.split('.')[0]
        section_name = "البناء الفكري" if section_number == "1" else "البناء اللغوي"

        # Create a more conversational prompt that includes the user's requirements
        prompt = f"""
        You are an expert Arabic language teacher chatbot. I have an Arabic exam with a reading passage about a specific theme.
        I want you to regenerate only the question {question_id} in the {section_name} section.

        Here is the reading passage from the exam:

        {reading_passage}

        CURRENT FULL EXAM:
        {exam_part}

        CURRENT CORRECTION:
        {correction_part}

        USER'S SPECIFIC REQUIREMENTS FOR THE NEW QUESTION:
        {user_requirements if user_requirements else "Create a different question that fits well with the reading passage."}

        TASK:
        1. Generate ONLY a new question for {question_id} in the {section_name} section that meets the user's requirements.
        2. Generate ONLY the matching answer for the correction sheet.
        3. Do not change any other parts of the exam or correction.
        4. For الادماجية الوضعية, keep it unchanged.
        5. If the question involves finding synonyms or antonyms, ensure the requested words appear verbatim in the passage.
        6. Maintain the appropriate difficulty level and style consistent with the rest of the exam.

        FORMAT YOUR RESPONSE EXACTLY LIKE THIS:

        NEW_QUESTION: [Your new question for {question_id} here]

        NEW_ANSWER: [Your new answer for {question_id} here]
        """

        try:
            response = client.models.generate_content(
                model=MODEL_ID,
                contents=prompt,
            )
            response_text = response.text

            # Extract new content
            question_match = re.search(r"NEW_QUESTION:(.*?)(?=NEW_ANSWER:|$)", response_text, re.DOTALL)
            answer_match = re.search(r"NEW_ANSWER:(.*?)(?=$)", response_text, re.DOTALL)

            if not question_match or not answer_match:
                print("❌ Error: Could not extract new question or answer.")
                return full_exam_text

            new_question = question_match.group(1).strip()
            new_answer = answer_match.group(1).strip()

            # Update exam parts without printing changes
            question_pattern = f"({re.escape(question_id)}\\s*\\.?\\s*)(.*?)(?=\\s*\\n\\s*(?:\\d+\\.\\d+|\\d+\\.|الوضعية|\\*\\*CORRECTION:\\*\\*|$))"
            
            updated_exam_part = re.sub(question_pattern, f"\\1{new_question}", exam_part, flags=re.DOTALL)
            updated_correction_part = re.sub(question_pattern, f"\\1{new_answer}", correction_part, flags=re.DOTALL)
            updated_full_exam = f"**EXAM:**\n{updated_exam_part}\n\n**CORRECTION:**\n{updated_correction_part}"
            
            # Save updated exam and show only success message
            with open("updated_arabic_exam.txt", "w", encoding="utf-8") as f:
                f.write(updated_full_exam)
            print("✅ Updated exam saved to 'updated_arabic_exam.txt'")

            return updated_full_exam

        except Exception as e:
            print(f"❌ Regeneration failed: {e}")
            return full_exam_text

def extract_reading_passage(exam_text):
    """Extract the reading passage from the exam text."""
    # Handle both marked and unmarked exam formats
    if "**EXAM:**" in exam_text:
        exam_part = exam_text.split("**CORRECTION:**")[0].replace("**EXAM:**", "").strip()
    else:
        exam_part = exam_text.split("**CORRECTION:**")[0].strip()

    # Pattern that better matches the actual output format
    match = re.search(r"النص\s*\n\s*(.*?)(?=\s*\n\s*1\.\s*البناء الفكري)", exam_part, re.DOTALL)
    if match:
        return match.group(1).strip()
    else:
        print("❌ Warning: Could not extract reading passage. Using empty passage.")
        return ""


def extract_exam_sections(full_exam_text, school_name, town_name, semester='1'):
    """Extract different sections from the exam text."""
    sections = {}

    # Split exam and correction parts
    if "**EXAM:**" in full_exam_text and "**CORRECTION:**" in full_exam_text:
        try:
            parts = full_exam_text.split("**CORRECTION:**", 1)
            exam_part = parts[0].replace("**EXAM:**", "").strip()
            correction_part = parts[1].strip()
        except IndexError:
            print("⚠️ Warning: Could not properly split exam and correction parts.")
            exam_part = full_exam_text
            correction_part = ""
    else:
        try:
            parts = full_exam_text.split("**CORRECTION:**", 1)
            if len(parts) != 2:
                print("⚠️ Warning: Could not split exam and correction parts.")
                exam_part = full_exam_text
                correction_part = ""
            else:
                exam_part = parts[0].strip()
                correction_part = parts[1].strip()
        except Exception as e:
            print(f"⚠️ Warning: Error splitting exam text: {e}")
            exam_part = full_exam_text
            correction_part = ""

    # Extract header with improved pattern matching
    header_patterns = [
        r"الجمهورية الجزائرية الديمقراطية الشعبية\s*(.*?)(?=\s*\n\s*النص)",
        r"الجمهورية الجزائرية\s*(.*?)(?=\s*\n\s*النص)",
        r"وزارة التربية\s*(.*?)(?=\s*\n\s*النص)"
    ]

    header_found = False
    for pattern in header_patterns:
        header_match = re.search(pattern, exam_part, re.DOTALL)
        if header_match:
            header_text = "الجمهورية الجزائرية الديمقراطية الشعبية\n" + header_match.group(1).strip()
            header_found = True
            break

    if not header_found:
        header_text = f"""الجمهورية الجزائرية الديمقراطية الشعبية
وزارة التربية الوطنية
مديرية التربية لولاية {town_name}
مدرسة {school_name}
اختبار الفصل {semester} في مادة اللغة العربية"""

    sections["header"] = header_text

    # Extract text passage with improved pattern matching
    text_patterns = [
        r"النص\s*\n\s*(.*?)(?=\s*\n\s*1\.\s*البناء الفكري)",
        r"النص\s*[:]\s*(.*?)(?=\s*\n\s*1\.\s*البناء)",
        r"النص\s*(.*?)(?=\s*\n\s*البناء الفكري)"
    ]

    text_found = False
    for pattern in text_patterns:
        text_match = re.search(pattern, exam_part, re.DOTALL)
        if text_match:
            sections["text"] = text_match.group(1).strip()  # Remove escape_latex_special_chars
            text_found = True
            break

    if not text_found:
        sections["text"] = ""
        print("⚠️ Warning: Could not extract reading passage.")

    # Extract البناء الفكري with improved pattern matching
    fikri_patterns = [
        r"1\.\s*البناء الفكري\s*\n(.*?)(?=\s*\n\s*2\.\s*البناء اللغوي)",
        r"البناء الفكري\s*\n(.*?)(?=\s*\n\s*البناء اللغوي)",
        r"1[\s\.:]*البناء الفكري\s*(.*?)(?=\s*\n\s*2[\s\.:]*البناء)"
    ]

    fikri_found = False
    for pattern in fikri_patterns:
        fikri_match = re.search(pattern, exam_part, re.DOTALL)
        if fikri_match:
            sections["fikri"] = fikri_match.group(1).strip()  # Remove escape_latex_special_chars
            fikri_found = True
            break

    if not fikri_found:
        sections["fikri"] = ""
        print("⚠️ Warning: Could not extract البناء الفكري section.")

    # Extract البناء اللغوي with improved pattern matching
    lughawi_patterns = [
        r"2\.\s*البناء اللغوي\s*\n(.*?)(?=\s*\n\s*الوضعية الادماجية)",
        r"البناء اللغوي\s*\n(.*?)(?=\s*\n\s*الوضعية)",
        r"2[\s\.:]*البناء اللغوي\s*(.*?)(?=\s*\n\s*الوضعية)"
    ]

    lughawi_found = False
    for pattern in lughawi_patterns:
        lughawi_match = re.search(pattern, exam_part, re.DOTALL)
        if lughawi_match:
            sections["lughawi"] = lughawi_match.group(1).strip()  # Remove escape_latex_special_chars
            lughawi_found = True
            break

    if not lughawi_found:
        sections["lughawi"] = ""
        print("⚠️ Warning: Could not extract البناء اللغوي section.")

    # Extract الوضعية الادماجية with improved pattern matching
    idmajiya_patterns = [
        r"الوضعية الادماجية\s*\n\s*(.*?)$",
        r"الوضعية الإدماجية\s*\n\s*(.*?)$",
        r"الوضعية\s*(.*?)$"
    ]

    idmajiya_found = False
    for pattern in idmajiya_patterns:
        idmajiya_match = re.search(pattern, exam_part, re.DOTALL)
        if idmajiya_match:
            sections["idmajiya"] = idmajiya_match.group(1).strip()  # Remove escape_latex_special_chars
            idmajiya_found = True
            break

    if not idmajiya_found:
        sections["idmajiya"] = ""
        print("⚠️ Warning: Could not extract الوضعية الادماجية section.")

    # Extract correction sections with similar improved pattern matching
    # Similar pattern improvements for the correction sections...
    # البناء الفكري (correction)
    for pattern in fikri_patterns:
        fikri_corr_match = re.search(pattern, correction_part, re.DOTALL)
        if fikri_corr_match:
            sections["fikri_corr"] = fikri_corr_match.group(1).strip()
            break
    else:
        sections["fikri_corr"] = ""
        print("⚠️ Warning: Could not extract البناء الفكري correction section.")

    # البناء اللغوي (correction)
    for pattern in lughawi_patterns:
        lughawi_corr_match = re.search(pattern, correction_part, re.DOTALL)
        if lughawi_corr_match:
            sections["lughawi_corr"] = lughawi_corr_match.group(1).strip()
            break
    else:
        sections["lughawi_corr"] = ""
        print("⚠️ Warning: Could not extract البناء اللغوي correction section.")

    # الوضعية الادماجية (correction)
    for pattern in idmajiya_patterns:
        idmajiya_corr_match = re.search(pattern, correction_part, re.DOTALL)
        if idmajiya_corr_match:
            sections["idmajiya_corr"] = idmajiya_corr_match.group(1).strip()
            break
    else:
        sections["idmajiya_corr"] = ""
        print("⚠️ Warning: Could not extract الوضعية الادماجية correction section.")

    # Debug information
    missing_sections = [key for key, value in sections.items() if not value]
    if missing_sections:
        print(f"⚠️ Warning: Empty sections detected: {', '.join(missing_sections)}")

    return sections

def generate_word_document(sections, output_path="arabic_exam.docx"):
    """Generate Word document directly using python-docx."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document()
    
    # Set RTL direction and font for the entire document
    style = doc.styles['Normal']
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style.font.rtl = True
    style.font.name = 'Arial'  # Use a common font that supports Arabic
    style.font.size = Pt(11)
    
    # Header
    for line in sections["header"].split('\n'):
        if line.strip():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(14)
            run.font.rtl = True
    
    doc.add_paragraph()  # Spacing
    
    # Reading passage
    heading = doc.add_heading('النص', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para = doc.add_paragraph(sections["text"])
    para.style.font.rtl = True
    
    doc.add_paragraph()  # Spacing
    
    # Questions sections
    def add_questions_section(title, content):
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for line in content.split('\n'):
            if line.strip():
                para = doc.add_paragraph()
                if re.match(r'^\d+\.\d+\s', line):
                    # Question ID in bold
                    question_id, rest = re.match(r'^(\d+\.\d+)\s(.+)$', line).groups()
                    run = para.add_run(question_id + ' ')
                    run.bold = True
                    run.font.rtl = True
                    run = para.add_run(rest)
                    run.font.rtl = True
                else:
                    run = para.add_run(line)
                    run.font.rtl = True
    
    # Add البناء الفكري section
    add_questions_section('1.البناء الفكري', sections["fikri"])
    doc.add_paragraph()
    
    # Add البناء اللغوي section
    add_questions_section('2.البناء اللغوي', sections["lughawi"])
    doc.add_paragraph()
    
    # Add الوضعية الادماجية section
    heading = doc.add_heading('الوضعية الادماجية', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para = doc.add_paragraph(sections["idmajiya"])
    para.style.font.rtl = True
    
    # Add page break before correction
    doc.add_page_break()
    
    # Add correction title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('التصحيح النموذجي للاختبار')
    run.bold = True
    run.font.size = Pt(14)
    run.font.rtl = True
    
    doc.add_paragraph()
    
    # Add correction sections
    add_questions_section('1.البناء الفكري', sections["fikri_corr"])
    doc.add_paragraph()
    add_questions_section('2.البناء اللغوي', sections["lughawi_corr"])
    doc.add_paragraph()
    
    # Add الوضعية الادماجية correction
    heading = doc.add_heading('الوضعية الادماجية', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para = doc.add_paragraph(sections["idmajiya_corr"])
    para.style.font.rtl = True
    
    try:
        doc.save(output_path)
        print(f"✅ Word document saved to '{output_path}'")
        return True
    except Exception as e:
        print(f"❌ Error saving Word document: {e}")
        return False

def run_chatbot():
    """Run the chatbot interface for the Arabic exam generator."""
    print("=" * 60)
    print("🤖 Welcome to the Arabic Exam Generator Chatbot! 🤖")
    print("=" * 60)
    print("\nThis chatbot will help you create and customize Arabic language exams.")

    # Prompt user for grade level
    while True:
        grade = input("\n💬 Which grade level do you want to create an exam for? (2, 3, 4, or 5): ")
        if grade in ["2", "3", "4", "5"]:
            break
        print("❌ Invalid grade. Please choose 2, 3, 4, or 5.")

    # Prompt user for semester
    while True:
        semester = input("\n💬 Which semester is this exam for? (1, 2, or 3): ")
        if semester in ["1", "2", "3"]:
            break
        print("❌ Invalid semester. Please choose 1, 2, or 3.")

    # Prompt user for exam theme
    user_theme = input("\n💬 What theme/topic would you like for the Arabic exam reading passage? ")
    if not user_theme.strip():
        print("No theme entered, using a generic topic.")
        user_theme = "موضوع عام"

    # Prompt user for school name and town
    school_name = input("\n💬 What's the name of the school? ")
    if not school_name.strip():
        print("No school name entered, using placeholder.")
        school_name = "..."

    town_name = input("\n💬 What's the town/province name? ")
    if not town_name.strip():
        print("No town name entered, using placeholder.")
        town_name = "..."

    print("\n🔄 Generating your initial Arabic exam... Please wait...")

    # Get the correct exam folders based on grade and semester
    exam_folders = get_exam_folders(grade, semester)

    # Generate initial exam with school and town information
    full_exam_text = generate_initial_exam(user_theme, school_name, town_name, grade, semester, exam_folders)
    if not full_exam_text:
        print("❌ Failed to generate initial exam.")
        return

    # Rest of the function remains unchanged...
    # Extract reading passage for future use
    reading_passage = extract_reading_passage(full_exam_text)

    # Store the current version of the exam for LaTeX generation
    current_exam_text = full_exam_text

    # Allow the user to interact with the chatbot
    while True:
        print("\n" + "=" * 60)
        print("🤖 Arabic Exam Chatbot - What would you like to do? 🤖")
        print("=" * 60)
        print("1. Edit a specific question (chat mode)")
        print("2. Generate Word document")
        print("3. Exit")
        choice = input("\n💬 Please enter your choice (1-3): ")

        if choice == "1":
            print("\n💬 What would you like to edit?")
            print("1. A specific question in البناء الفكري or البناء اللغوي")
            print("2. The integration situation (الوضعية الادماجية)")
            edit_choice = input("\n💬 Please enter your choice (1-2): ")

            if edit_choice == "1":
                question_id = input("\n💬 Which question would you like to change? (e.g., 1.1, 2.3): ")
                # Validate question ID format
                if not re.match(r"^[12]\.[1-9]$", question_id):
                    print("❌ Invalid question ID format. Use format like 1.1 or 2.3")
                    continue

                section_name = "البناء الفكري (conceptual construction)" if question_id.startswith("1") else "البناء اللغوي (linguistic construction)"
                print(f"\n💬 You're editing question {question_id} in the {section_name} section.")

            elif edit_choice == "2":
                question_id = "idmajiya"
                print("\n💬 You're editing the integration situation (الوضعية الادماجية) section.")

            else:
                print("❌ Invalid choice. Please enter 1 or 2.")
                continue

            # Get user's detailed requirements in a chat-like manner
            print("\n💬 Please describe in detail how you want this changed. You can mention:")
            if edit_choice == "1":
                print("   - The type of question you want (e.g., synonym, antonym, grammar, etc.)")
                print("   - The difficulty level")
                print("   - Any specific words or grammatical concepts to include")
            else:
                print("   - The type of integration situation you want")
                print("   - The skills you want students to demonstrate")
                print("   - Any specific theme or focus you want to include")

            print("   - Any other preferences you have")
            user_requirements = input("\n🔤 Your detailed instructions: ")

            print(f"\n🔄 Regenerating {question_id if question_id != 'idmajiya' else 'الوضعية الادماجية'} based on your requirements... Please wait...")

            # Regenerate the specific question or integration situation with user's requirements
            updated_exam_text = regenerate_specific_question(current_exam_text, question_id, reading_passage, user_requirements)

            # Update the current exam text for future operations
            current_exam_text = updated_exam_text

            # Re-extract the reading passage (in case it was modified)
            reading_passage = extract_reading_passage(current_exam_text)

            print("\n✅ Content updated according to your specifications!")
            print("\n💬 Would you like to make any other changes?")

        elif choice == "2":
            print("\n🔄 Generating Word document... Please wait...")
            
            # Extract sections from the current exam text
            sections = extract_exam_sections(current_exam_text, school_name, town_name, semester)
            if not sections:
                print("❌ Error extracting exam sections for Word formatting.")
                continue
            
            # Generate Word document
            docx_path = "arabic_exam.docx"
            if generate_word_document(sections, docx_path):
                print(f"\n✅ Word document saved to '{docx_path}'")
            else:
                print("❌ Error generating Word document.")

        elif choice == "3":
            print("\n👋 Thank you for using the Arabic Exam Generator Chatbot! Goodbye!")
            break

        else:
            print("❌ Invalid choice. Please enter a number between 1 and 3.")

# 5. Update the main function to call the modified run_chatbot
def main():
    """Main function to run the chatbot interface."""
    try:
        run_chatbot()
    except KeyboardInterrupt:
        print("\n\n👋 Chatbot session interrupted. Goodbye!")
    except Exception as e:
        print(f"\n❌ An unexpected error occurred: {e}")
        print("Please try again or contact support.")

if __name__ == "__main__":
    main()